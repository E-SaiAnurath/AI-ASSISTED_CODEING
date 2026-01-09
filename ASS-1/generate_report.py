"""
Generate a Word report `Report.docx` summarizing Tasks 1-5: prompts, code,
explanations, and sample outputs. Uses `python-docx`.
"""
import runpy
import io
import sys
from pathlib import Path
from time import perf_counter

try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    raise SystemExit("python-docx is required. Run: pip install python-docx")


def read_file(path):
    p = Path(path)
    return p.read_text(encoding='utf-8') if p.exists() else ''


def capture_task5_outputs():
    module = runpy.run_path('Task5.py')
    rl = module['reverse_loop']
    rs = module['reverse_slice']
    samples = ["Hello World", "abc123", "" ]
    buf = io.StringIO()
    for s in samples:
        buf.write(f"Input: {s!r}\n")
        buf.write(f"  loop: {rl(s)!r}\n")
        buf.write(f"  slice: {rs(s)!r}\n")
    # performance test on large input
    big = 'a' * 1000000
    t0 = perf_counter(); rl(big); t1 = perf_counter()
    t_loop = t1 - t0
    t0 = perf_counter(); rs(big); t1 = perf_counter()
    t_slice = t1 - t0
    buf.write(f"\nLarge-input timings (1e6 chars): loop={t_loop:.4f}s slice={t_slice:.4f}s\n")
    return buf.getvalue()


def build_doc():
    doc = Document()
    doc.add_heading('AI-generated tasks report', level=1)

    # Include files
    for name in ('Task1.py', 'Task2.py', 'Task3.py', 'Task5.py'):
        doc.add_heading(name, level=2)
        code = read_file(name)
        p = doc.add_paragraph()
        run = p.add_run(code)
        font = run.font
        font.name = 'Courier New'
        font.size = Pt(9)

    doc.add_heading('Task5 Outputs & Performance', level=2)
    out = capture_task5_outputs()
    doc.add_paragraph(out)

    doc.add_heading('Comparative Notes', level=2)
    notes = (
        "Loop vs slicing: both O(n) time and O(n) memory for creating reversed string.\n"
        "Slicing is implemented in C and has lower constant overhead; loop gives more control for custom logic.\n"
        "Use slicing for simplicity; use loop when you need to transform elements during reversal."
    )
    doc.add_paragraph(notes)

    out_path = Path('Report.docx')
    doc.save(str(out_path))
    return out_path


if __name__ == '__main__':
    print('Generating Report.docx...')
    p = build_doc()
    print('Saved', p)
